import ollama
from docx import Document
import json
import subprocess
import time

# ======================================
# 1. CARGAR Y VALIDAR ARCHIVOS
# ======================================
def safe_read_docx(file_path):
    try:
        doc = Document(file_path)
        if len(doc.paragraphs) == 0:
            raise ValueError("Documento vacío o inválido")
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise RuntimeError(f"Error leyendo DOCX: {str(e)}")

try:
    story_text = safe_read_docx("documento.docx")
    print(f"\n✅ Documento cargado - Caracteres: {len(story_text)}")

    with open("entrenamiento.json") as f:
        training_data = json.load(f)
    print(f"✅ JSON cargado - Ejemplos: {len(training_data)}")
except Exception as e:
    print(f"\n🔥 Error crítico: {str(e)}")
    raise

# ======================================
# 2. CONFIGURAR OLLAMA Y MODELO
# ======================================
modelfile_content = f'''
FROM llama3
SYSTEM "Eres un asistente especializado en analizar la historia del usuario. Contexto clave: {story_text[:1000]}"
'''

with open("Modelfile", "w") as f:
    f.write(modelfile_content)

print("\n🛠️ Creando modelo personalizado...")
try:
    subprocess.run(["ollama", "create", "mi-historia", "-f", "Modelfile"], check=True)
    subprocess.run(["ollama", "list"], check=True)
except subprocess.CalledProcessError as e:
    print(f"\n🔥 Error al crear el modelo: {str(e)}")
    raise

# ======================================
# 3. ENTRENAMIENTO CON MANEJO DE ERRORES
# ======================================
def train_model():
    formatted_data = []
    for item in training_data:
        formatted_data.append({
            "instruction": item["input"],
            "context": story_text[:1500],
            "response": item["output"]
        })

    for epoch in range(3):
        print(f"\n🏋️ Época {epoch+1}/3")
        for i, example in enumerate(formatted_data[:50]):
            try:
                response = ollama.generate(
                    model='mi-historia',
                    prompt=f"Pregunta: {example['instruction']}",
                    system=f"Contexto: {example['context']}\nRespuesta esperada: {example['response']}",
                    options={'num_ctx': 2048, 'temperature': 0.4}
                )
                if i % 5 == 0:
                    print(f"✅ Progreso: {i+1}/{len(formatted_data)}")
            except Exception as e:
                print(f"⚠️ Error en ejemplo {i}: {str(e)}...")
                time.sleep(5)

train_model()

# ======================================
# 4. PRUEBA FINAL DEL MODELO
# ======================================
def ask_question(pregunta):
    try:
        response = ollama.generate(
            model='mi-historia',
            prompt=f"Pregunta: {pregunta}",
            system=f"Contexto completo: {story_text[:1500]}",
            options={'temperature': 0.5}
        )
        return response['response']
    except Exception as e:
        return f"Error: {str(e)}"

print("\n🧪 Prueba del modelo entrenado:")
pregunta_test = "¿Qué impacto tuvo la plataforma de aprendizaje en los jóvenes del barrio?"
print(f"Pregunta: {pregunta_test}")
print(f"Respuesta: {ask_question(pregunta_test)}")
